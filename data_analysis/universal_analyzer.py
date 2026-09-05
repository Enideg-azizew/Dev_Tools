#!/usr/bin/env python3
"""
Universal Statistical Analyzer.
Combines analyzer.py and analyzer1.py.
"""
import sys
import warnings
from pathlib import Path
from typing import List, Dict, Any, Optional
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
from scipy.stats import shapiro, levene
import statsmodels.api as sm
from statsmodels.formula.api import logit, ols
from docx import Document
from docx.shared import Inches
from datetime import datetime

warnings.filterwarnings('ignore')
plt.style.use('seaborn-v0_8-whitegrid')
sns.set_palette("husl")

class UniversalAnalyzer:
    """Comprehensive statistical analysis tool."""
    
    def __init__(self):
        self.df = None
        self.results = {}
        self.config = {}
        self.variable_types = {}
    
    def load_data(self, file_path: str, sheet_name: Optional[str] = None) -> pd.DataFrame:
        """Load data from various formats."""
        path = Path(file_path)
        
        if path.suffix in ['.xlsx', '.xls']:
            self.df = pd.read_excel(file_path, sheet_name=sheet_name)
            if isinstance(self.df, dict):
                self.df = self.df[list(self.df.keys())[0]]
        elif path.suffix == '.csv':
            self.df = pd.read_csv(file_path)
        elif path.suffix == '.json':
            self.df = pd.read_json(file_path)
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")
        
        self._detect_variable_types()
        return self.df
    
    def _detect_variable_types(self) -> None:
        """Auto-detect variable types."""
        self.variable_types = {
            'categorical': [],
            'continuous': [],
            'binary': [],
            'ordinal': [],
            'datetime': []
        }
        
        for col in self.df.columns:
            if self.df[col].isnull().mean() > 0.5:
                continue
            
            if pd.api.types.is_datetime64_any_dtype(self.df[col]):
                self.variable_types['datetime'].append(col)
            elif self.df[col].nunique() <= 10:
                if self.df[col].nunique() == 2:
                    self.variable_types['binary'].append(col)
                else:
                    self.variable_types['categorical'].append(col)
            elif pd.api.types.is_string_dtype(self.df[col]) and self.df[col].nunique() <= 20:
                self.variable_types['ordinal'].append(col)
            else:
                self.variable_types['continuous'].append(col)
    
    def set_config(self, outcome_var: Optional[str] = None,
                   group_var: Optional[str] = None,
                   covariates: Optional[List[str]] = None,
                   alpha: float = 0.05) -> None:
        """Set analysis configuration."""
        self.config = {
            'outcome_var': outcome_var,
            'group_var': group_var,
            'covariates': covariates or [],
            'alpha': alpha
        }
    
    def data_quality_report(self) -> Dict[str, Any]:
        """Generate data quality report."""
        return {
            'shape': self.df.shape,
            'missing': self.df.isnull().sum().to_dict(),
            'missing_pct': (self.df.isnull().mean() * 100).to_dict(),
            'duplicates': self.df.duplicated().sum(),
            'data_types': self.df.dtypes.astype(str).to_dict()
        }
    
    def calculate_prevalence(self, outcome_var: Optional[str] = None) -> Dict[str, Any]:
        """Calculate prevalence for binary outcome."""
        var = outcome_var or self.config.get('outcome_var')
        if not var:
            raise ValueError("Outcome variable not specified")
        
        n_total = len(self.df)
        n_positive = self.df[var].sum()
        p = n_positive / n_total
        z = 1.96
        se = np.sqrt((p * (1 - p)) / n_total)
        
        return {
            'n_total': n_total,
            'n_positive': n_positive,
            'prevalence': p * 100,
            'ci_lower': (p - z * se) * 100,
            'ci_upper': (p + z * se) * 100
        }
    
    def bivariate_analysis(self, outcome_var: Optional[str] = None) -> List[Dict[str, Any]]:
        """Perform bivariate analysis."""
        var = outcome_var or self.config.get('outcome_var')
        if not var:
            raise ValueError("Outcome variable not specified")
        
        results = []
        
        for col in self.df.columns:
            if col == var:
                continue
            
            if self.df[col].isnull().mean() > 0.3:
                continue
            
            if col in self.variable_types['continuous']:
                if self.df[var].nunique() == 2:
                    group0 = self.df[self.df[var] == 0][col].dropna()
                    group1 = self.df[self.df[var] == 1][col].dropna()
                    
                    if len(group0) > 1 and len(group1) > 1:
                        # Check assumptions
                        _, p0 = shapiro(group0) if len(group0) <= 5000 else (0, 0.05)
                        _, p1 = shapiro(group1) if len(group1) <= 5000 else (0, 0.05)
                        _, p_levene = levene(group0, group1)
                        
                        if p0 > 0.05 and p1 > 0.05 and p_levene > 0.05:
                            stat, p_value = stats.ttest_ind(group0, group1)
                            test = 't-test'
                        else:
                            stat, p_value = stats.mannwhitneyu(group0, group1)
                            test = 'Mann-Whitney U'
                        
                        results.append({
                            'variable': col,
                            'type': 'continuous',
                            'test': test,
                            'statistic': stat,
                            'p_value': p_value
                        })
            
            elif col in self.variable_types['categorical'] or col in self.variable_types['binary']:
                contingency = pd.crosstab(self.df[col], self.df[var])
                if contingency.shape[0] > 1 and contingency.shape[1] > 1:
                    chi2, p_value, dof, expected = stats.chi2_contingency(contingency)
                    results.append({
                        'variable': col,
                        'type': 'categorical',
                        'test': 'Chi-square',
                        'statistic': chi2,
                        'p_value': p_value,
                        'degrees_of_freedom': dof
                    })
        
        return sorted(results, key=lambda x: x.get('p_value', 1))
    
    def multivariate_analysis(self, method: str = 'logistic') -> Optional[Dict[str, Any]]:
        """Perform multivariate analysis."""
        outcome_var = self.config.get('outcome_var')
        if not outcome_var:
            return None
        
        # Get candidate variables from bivariate analysis
        bivariate = self.bivariate_analysis()
        candidates = [r['variable'] for r in bivariate if r.get('p_value', 1) < 0.2]
        
        if not candidates:
            return None
        
        formula = f"{outcome_var} ~ " + " + ".join(candidates)
        
        try:
            if method == 'logistic' and self.df[outcome_var].nunique() == 2:
                model = logit(formula, data=self.df).fit(disp=False)
            else:
                model = ols(formula, data=self.df).fit()
            
            return {
                'formula': formula,
                'method': method,
                'summary': model.summary(),
                'params': model.params.to_dict(),
                'pvalues': model.pvalues.to_dict(),
                'r_squared': model.rsquared if hasattr(model, 'rsquared') else None
            }
        except Exception as e:
            return {'error': str(e)}
    
    def create_report(self, output_path: str = 'analysis_report.docx') -> str:
        """Generate Word report."""
        doc = Document()
        doc.add_heading('Statistical Analysis Report', 0)
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph(f'Dataset: {len(self.df)} rows, {len(self.df.columns)} columns')
        
        # Quality
        doc.add_heading('Data Quality', level=1)
        quality = self.data_quality_report()
        for col, missing in quality['missing'].items():
            if missing > 0:
                doc.add_paragraph(f"{col}: {missing} missing ({quality['missing_pct'][col]:.1f}%)")
        
        doc.save(output_path)
        return output_path

def main():
    if len(sys.argv) < 2:
        print("Usage: python universal_analyzer.py <data_file> [outcome_var]")
        sys.exit(1)
    
    analyzer = UniversalAnalyzer()
    analyzer.load_data(sys.argv[1])
    
    if len(sys.argv) > 2:
        analyzer.set_config(outcome_var=sys.argv[2])
        results = analyzer.bivariate_analysis()
        print("\nBivariate Analysis Results:")
        for r in results[:10]:
            print(f"{r['variable']}: p={r.get('p_value', 'N/A'):.4f}")
    
    analyzer.create_report()

if __name__ == "__main__":
    main()
